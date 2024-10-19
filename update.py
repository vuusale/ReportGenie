from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.pdfgen import canvas
import os

def update_table_of_contents(docx_file_path):
    try:
        # Load the document
        doc = Document(docx_file_path)

        # Find and update the TOC field
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if 'TOC' in run.text:
                    fldChar = OxmlElement('w:fldChar')
                    fldChar.set(qn('w:fldCharType'), 'begin')

                    instrText = OxmlElement('w:instrText')
                    instrText.set(qn('xml:space'), 'preserve')
                    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

                    fldChar2 = OxmlElement('w:fldChar')
                    fldChar2.set(qn('w:fldCharType'), 'separate')

                    fldChar3 = OxmlElement('w:fldChar')
                    fldChar3.set(qn('w:fldCharType'), 'end')

                    r_element = run._r
                    r_element.append(fldChar)
                    r_element.append(instrText)
                    r_element.append(fldChar2)
                    r_element.append(fldChar3)

        # Save the updated document
        doc.save(docx_file_path)
        print(f"Table of contents updated successfully for: {docx_file_path}")

    except Exception as e:
        print(f"An error occurred while updating the table of contents: {e}")

def convert_docx_to_pdf(docx_file_path, pdf_file_path):
    try:
        # Load the DOCX file
        doc = Document(docx_file_path)

        # Create a PDF canvas
        c = canvas.Canvas(pdf_file_path)

        # Set starting position
        x, y = 40, 800

        # Iterate over paragraphs in DOCX and add them to the PDF
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if text:
                c.drawString(x, y, text)
                y -= 20
                if y < 40:  # Start a new page if at the bottom
                    c.showPage()
                    y = 800

        # Save the PDF
        c.save()
    
    except Exception as e:
        print(f"An error occurred while converting DOCX to PDF: {e}")

# Usage Example
docx_file_path = 'pentest_report_output.docx'  # Path to the docx file to be updated
update_table_of_contents(docx_file_path)
