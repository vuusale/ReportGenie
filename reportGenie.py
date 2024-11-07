from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import io
from docx.shared import Inches
from datetime import datetime
from docx.opc.part import Part
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import pypandoc
import io
import tempfile
import os
from utils import * 

def generate_pentest_report(report_title, start_date, end_date, reporter_name, vulnerabilities, icon_path, executive_summary, output_path):
    template_path = 'report.docx'
    doc = Document(template_path)
    start_date = datetime(*[int(i) for i in start_date.split("-")])
    end_date = datetime(*[int(i) for i in end_date.split("-")])
    date = format_date_range(start_date, end_date)
    for paragraph in doc.paragraphs:
        if '{REPORT_TITLE}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{REPORT_TITLE}', report_title)
        if '{DATE}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{DATE}', date)
        if '{REPORTER_NAME}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{REPORTER_NAME}', reporter_name)

    for paragraph in doc.paragraphs:
        if '{ICON}' in paragraph.text:
            paragraph.text = ''
            run = paragraph.add_run()
            run.add_picture(icon_path, height=Inches(1.2))
            run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(18)
            break

    for paragraph in doc.paragraphs:
        if '{TOC}' in paragraph.text:
            paragraph.clear()
            add_toc(doc, paragraph)
            break

    for i, vuln in enumerate(vulnerabilities, start=1):
        heading = doc.add_heading(level=2)
        heading_run = heading.add_run(f"{i}. {vuln.vulnerability_name}")

        heading.paragraph_format.line_spacing = 1.5

        severity_paragraph = doc.add_paragraph()
        severity_run = severity_paragraph.add_run("Severity: ")
        severity_run.bold = True
        severity_value_run = severity_paragraph.add_run(vuln.severity)
        severity_value_run.bold = True
        severity_value = vuln.severity
        if severity_value.lower() == 'critical':
            severity_value_run.font.color.rgb = RGBColor(128, 0, 0)
        elif severity_value.lower() == 'high':
            severity_value_run.font.color.rgb = RGBColor(255, 0, 0)
        elif severity_value.lower() == 'medium':
            severity_value_run.font.color.rgb = RGBColor(255, 165, 0)
        elif severity_value.lower() == 'low':
            severity_value_run.font.color.rgb = RGBColor(0, 0, 255)
        elif severity_value.lower() == 'informational':
            severity_value_run.font.color.rgb = RGBColor(0, 128, 0)
        else:
            severity_value_run.font.color.rgb = RGBColor(0, 0, 0)

        severity_paragraph.paragraph_format.line_spacing = 1.5

        vulnerable_component_paragraph = doc.add_paragraph()
        vulnerable_component_run = vulnerable_component_paragraph.add_run("URL/Vulnerable component: ")
        vulnerable_component_run.bold = True
        vulnerable_component_value_run = vulnerable_component_paragraph.add_run(vuln.vulnerable_component)
        vulnerable_component_paragraph.paragraph_format.line_spacing = 1.5

        description_paragraph = doc.add_paragraph()
        description_run = description_paragraph.add_run("Description: ")
        description_run.bold = True
        add_html_after_paragraph(doc, "Description", vuln.description, description_paragraph)
        # add_html_after_paragraph(doc, 'Description', vuln.description, description_paragraph)

        impact_paragraph = doc.add_paragraph()
        impact_paragraph.paragraph_format.space_before = Pt(8)
        impact_run = impact_paragraph.add_run("Impact: ")
        impact_run.bold = True
        add_html_after_paragraph(doc, 'Impact', vuln.impact, impact_paragraph)

        remediation_paragraph = doc.add_paragraph()
        remediation_paragraph.paragraph_format.space_before = Pt(8)
        remediation_run = remediation_paragraph.add_run("Remediation: ")
        remediation_run.bold = True
        add_html_after_paragraph(doc, 'Remediation', vuln.remediation, remediation_paragraph)

        poc_paragraph = doc.add_paragraph()
        poc_paragraph.paragraph_format.space_before = Pt(8)
        poc_run = poc_paragraph.add_run("PoC: ")
        poc_run.bold = True
        add_html_after_paragraph(doc, 'PoC', vuln.poc, poc_paragraph)

    severity_counts = {
        'Critical': 0,
        'High': 0,
        'Medium': 0,
        'Low': 0,
        'Informational': 0,
        'Other': 0
    }
    
    for vuln in vulnerabilities:
        severity = severity_value.capitalize()
        if severity in severity_counts:
            severity_counts[severity] += 1
        else:
            severity_counts['Other'] += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '{CRITICAL}' in cell.text:
                    cell.text = cell.text.replace('{CRITICAL}', str(severity_counts['Critical']))
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if '{HIGH}' in cell.text:
                    cell.text = cell.text.replace('{HIGH}', str(severity_counts['High']))
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if '{MEDIUM}' in cell.text:
                    cell.text = cell.text.replace('{MEDIUM}', str(severity_counts['Medium']))
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if '{LOW}' in cell.text:
                    cell.text = cell.text.replace('{LOW}', str(severity_counts['Low']))
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if '{INFORMATIONAL}' in cell.text:
                    cell.text = cell.text.replace('{INFORMATIONAL}', str(severity_counts['Informational']))
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if '{TOTAL}' in cell.text:
                    cell.text = cell.text.replace('{TOTAL}', str(sum(severity_counts.values())))
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    labels = [label for label, count in severity_counts.items() if count > 0]
    sizes = [count for count in severity_counts.values() if count > 0]
    colors = ['#800000', '#FF0000', '#FFA500', '#0000FF', '#008000', '#800080'][:len(labels)]
    
    plt.switch_backend('Agg')
    plt.figure(figsize=(6, 6))
    plt.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
    plt.axis('equal')
    plt.title('Vulnerability Severity Distribution\n')
    
    pie_chart_stream = io.BytesIO()
    plt.savefig(pie_chart_stream, format='png')
    pie_chart_stream.seek(0)

    for paragraph in doc.paragraphs:
        if '{TECHNICAL_SUMMARY_CHART}' in paragraph.text:
            paragraph.text = ''
            run = paragraph.add_run()
            run.add_picture(pie_chart_stream, width=Inches(4.5))
            run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            break

    for paragraph in doc.paragraphs:
        if '{EXECUTIVE_SUMMARY}' in paragraph.text:
            paragraph.text = executive_summary
            break

    # Remove empty pages by iterating over paragraphs and deleting page breaks if followed by empty paragraphs
    # remove_empty_pages(doc)

    for section in doc.sections:
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = ''
        header_run = header_paragraph.add_run()
        header_run.add_picture(icon_path, height=Inches(0.4))
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_paragraph.paragraph_format.line_spacing = 1.5

    for paragraph in doc.paragraphs:
        if '{REPORT_TITLE}' in paragraph.text or '{DATE}' in paragraph.text or '{REPORTER_NAME}' in paragraph.text:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph.paragraph_format.space_before = Pt(600)
    
    doc.save(output_path)
    print(f"Report generated: {output_path}")

vulnerabilities = [
    {
        'vulnerability_name': 'SQL Injection',
        'vulnerable_component': 'http://example.com/login',
        'severity': 'High',
        'description': 'SQL injection vulnerability in login form.',
        'impact': 'Access to all user data.',
        'remediation': 'Use parameterized queries.',
        'poc': 'Use this to bypass authentication.'
    },
    {
        'vulnerability_name': 'Cross-Site Scripting',
        'vulnerable_component': 'http://example.com/comment',
        'severity': 'Medium',
        'description': 'Reflected XSS in comment section.',
        'impact': 'Stealing session cookies.',
        'remediation': 'Sanitize user input.',
        'poc': '<script>alert("XSS")</script>'
    },
    {
        'vulnerability_name': 'SSTI',
        'vulnerable_component': 'http://example.com/comment',
        'severity': 'Critical',
        'description': 'SSTI in order placing.',
        'impact': 'Remote Code Execution.',
        'remediation': 'Sanitize user input.',
        'poc': '{{7*7}}'
    }
]
vulnerabilities2 = [
    {
        'vulnerability_name': 'a',
        'vulnerable_component': 'http://example.com/login',
        'severity': 'High',
        'description': 'SQL injection vulnerability in login form.',
        'impact': 'Access to all user data.',
        'remediation': 'Use parameterized queries.',
        'poc': 'Use this to bypass authentication.'
    }
]

executive_summary = "a."
icon_path = 'logo.png'

if __name__ == "__main__":
    generate_pentest_report('Sample Pentest Report', '2024-01-01', 'John Doe', vulnerabilities2, icon_path, executive_summary)