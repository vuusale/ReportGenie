from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import io
from docx.shared import Inches
from datetime import datetime
from docx.opc.part import Part
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import pypandoc
import io
import tempfile
import os


def set_font_size_paragraph(paragraph, font, size):
    for run in paragraph.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(size)


def add_html_after_paragraph(doc, paragraph_text, html, paragraph):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_filename = temp_file.name
        pypandoc.convert_text(html, 'docx', format='html', outputfile=temp_filename)

    temp_doc = Document(temp_filename)

    for paragraph in doc.paragraphs:
        if paragraph_text in paragraph.text:

            current_paragraph_element = paragraph._element

            new_elements = []
            for temp_paragraph in temp_doc.paragraphs:
                set_font_size_paragraph(temp_paragraph, "Calibri", 14)
                new_paragraph_element = OxmlElement('w:p')
                new_paragraph_element.extend(temp_paragraph._element)
                new_elements.append(new_paragraph_element)

            for new_element in reversed(new_elements):
                current_paragraph_element.addnext(new_element)
            
            # new_paragraph = paragraph.insert_paragraph_before(temp_paragraph.text)
            # new_paragraph.paragraph_format.space_after = Pt(2)
            break

    os.remove(temp_filename)
    # set_font_size_paragraph(paragraph, "Calibri", 14)

def format_date_range(start_date, end_date):
    # Define date format
    start_fmt = "%-d %B"
    end_fmt = "%-d %B %Y"
    
    # Handle cases where months are the same or different
    if start_date.month == end_date.month:
        # If month is the same, format as "1-4 November 2024"
        formatted_range = f"{start_date.strftime('%-d')}-{end_date.strftime(end_fmt)}"
    else:
        # If month is different, format as "20 November-10 December 2024"
        formatted_range = f"{start_date.strftime(start_fmt)}-{end_date.strftime(end_fmt)}"
    
    return formatted_range

def remove_empty_pages(doc):
    for paragraph in doc.paragraphs:
        # Remove page breaks followed by empty paragraphs
        if paragraph._element.xpath('.//w:br[@w:type="page"]') and not paragraph.text.strip():
            parent = paragraph._element.getparent()
            parent.remove(paragraph._element)

def add_toc(doc, paragraph):
    run = paragraph.add_run("Table of Contents")
    run.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.bold = True
    run.underline = True

    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar3 = OxmlElement('w:updateFields') 
    fldChar3.set(qn('w:val'), 'true') 
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    # Automatically update TOC without requiring a manual update
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    doc.element.body.insert(0, update_fields)

    p_element = paragraph._p